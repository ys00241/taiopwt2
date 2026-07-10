"""Pre-event blueprint — 前期準備 (今年聖物, 支出, 贊助, 上年欠款, 粉紅紙, 收據)."""
import json
from io import BytesIO
from datetime import datetime
from zipfile import ZipFile

from flask import (
    Blueprint, render_template, request, jsonify, redirect, url_for, Response, send_file,
)
from flask_login import login_required

from app.extensions import db
from app.models.this_year_item import ThisYearItem
from app.models.expense import Expense
from app.models.sponsor import Sponsor
from app.models.member import Member
from app.models.bid import Bid
from app.models.item import Item

bp = Blueprint("pre_event", __name__,
               url_prefix="/pre")

# ──────────────────────────────────────────────────────────────────────
#  Expense categories (predefined)
# ──────────────────────────────────────────────────────────────────────
PREDEFINED_CATEGORIES = [
    "酒席", "場地", "佈置", "音響", "攝影",
    "神料", "花炮", "樂隊", "歌星", "舞獅",
    "工作人員", "雜項", "交通", "宣傳", "保險",
]


# ──────────────────────────────────────────────────────────────────────
#  A) 聖物準備 — this_year_items CRUD
# ──────────────────────────────────────────────────────────────────────

@bp.route("/items")
@login_required
def pre_items():
    """今年聖物列表 — Current year items with sticker numbers."""
    import logging
    logger = logging.getLogger(__name__)
    year = request.args.get("year", datetime.now().strftime("%Y"))
    year_int = int(year)
    items = (
        ThisYearItem.query
        .filter_by(year=year_int)
        .order_by(ThisYearItem.sticker_no)
        .all()
    )
    # Safety check: verify all returned items match the requested year
    mismatched = [i for i in items if i.year != year_int]
    if mismatched:
        logger.warning(
            "pre_items: %d item(s) with year mismatch (expected %s). IDs: %s",
            len(mismatched), year, [i.id for i in mismatched],
        )
    years = (
        db.session.query(ThisYearItem.year)
        .distinct()
        .order_by(ThisYearItem.year.desc())
        .all()
    )
    years = [r[0] for r in years]

    # Stats for stat cards
    total = len(items)
    unauctioned = sum(1 for i in items if not i.bidder_name)
    auctioned = total - unauctioned
    stats = {"total": total, "unauctioned": unauctioned, "auctioned": auctioned}

    items_json = json.dumps({str(i.id): dict(
        id=i.id,
        item_name=i.item_name,
        category=i.category,
        cost=i.cost,
        source=i.source,
        notes=i.notes,
        actual_item=i.actual_item,
        sticker_no=i.sticker_no,
    ) for i in items}, ensure_ascii=False)

    return render_template(
        "pre_event/items.html",
        items=items,
        year=year,
        years=years,
        stats=stats,
        items_json=items_json,
    )


@bp.route("/items/add", methods=["POST"])
@login_required
def pre_items_add():
    """Add a new this-year item with auto sticker assignment."""
    year = int(request.form.get("year", datetime.now().strftime("%Y")))
    item_name = request.form.get("item_name", "").strip()
    category = request.form.get("category", "").strip()
    cost = float(request.form.get("cost", 0) or 0)
    source = request.form.get("source", "").strip()
    notes = request.form.get("notes", "").strip()
    actual_item = request.form.get("actual_item", "").strip()
    photo_codes = request.form.get("photo_codes", "").strip()

    # Auto-assign sticker number: 1-50, skip 4,14,24,34
    used = {
        r[0] for r in db.session.query(ThisYearItem.sticker_no)
        .filter(ThisYearItem.year == year, ThisYearItem.sticker_no.isnot(None))
        .all()
    }
    skip = {4, 14, 24, 34}
    sticker_no = None
    if request.form.get("auto_sticker", "1") == "1":
        for n in range(1, 51):
            if n not in used and n not in skip:
                sticker_no = n
                break

    item = ThisYearItem(
        year=year,
        sticker_no=sticker_no,
        item_name=item_name,
        category=category,
        cost=cost,
        source=source,
        notes=notes,
        actual_item=actual_item,
        photo_codes=photo_codes,
    )
    db.session.add(item)
    db.session.commit()

    # If AJAX request, return JSON. Otherwise redirect back to items page.
    if request.headers.get("Accept") == "application/json" or request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return jsonify({"ok": True, "sticker_no": sticker_no, "item_id": item.id})
    return redirect(url_for("pre_event.pre_items"))


@bp.route("/items/<int:item_id>/edit", methods=["POST"])
@login_required
def pre_items_edit(item_id):
    """Edit a this-year item."""
    item = db.session.get(ThisYearItem, item_id)
    if not item:
        return jsonify({"error": "Item not found"}), 404

    item.item_name = request.form.get("item_name", item.item_name)
    item.category = request.form.get("category", item.category)
    item.cost = float(request.form.get("cost", 0) or 0)
    item.source = request.form.get("source", item.source)
    item.sticker_no = int(request.form.get("sticker_no", 0) or 0) or None
    item.notes = request.form.get("notes", item.notes)
    item.actual_item = request.form.get("actual_item", item.actual_item)
    item.photo_codes = request.form.get("photo_codes", item.photo_codes)

    # Also allow editing bid info from pre-event
    item.bidder_name = request.form.get("bidder_name", item.bidder_name)
    bid_amt = request.form.get("bid_amount")
    if bid_amt:
        item.bid_amount = float(bid_amt)

    db.session.commit()

    # If AJAX request, return JSON. Otherwise redirect back to items page.
    if request.headers.get("Accept") == "application/json" or request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return jsonify({"ok": True})
    return redirect(url_for("pre_event.pre_items"))


@bp.route("/items/<int:item_id>/delete", methods=["POST"])
@login_required
def pre_items_delete(item_id):
    """Delete a this-year item."""
    item = db.session.get(ThisYearItem, item_id)
    if not item:
        return jsonify({"error": "Item not found"}), 404
    db.session.delete(item)
    db.session.commit()
    return jsonify({"ok": True})


@bp.route("/items/import-csv", methods=["POST"])
@login_required
def pre_items_import_csv():
    """Import this_year_items from CSV with auto sticker assignment."""
    import csv as csv_mod
    import io

    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file"}), 400

    year = int(request.form.get("year", datetime.now().strftime("%Y")))
    content = f.read().decode("utf-8-sig")
    reader = csv_mod.DictReader(io.StringIO(content))

    # Get used sticker numbers for this year
    used = {
        r[0] for r in db.session.query(ThisYearItem.sticker_no)
        .filter(ThisYearItem.year == year, ThisYearItem.sticker_no.isnot(None))
        .all()
    }
    skip = {4, 14, 24, 34}
    next_sticker = 1

    count = 0
    for row in reader:
        item_name = row.get("item_name", "").strip()
        if not item_name:
            continue
        # Auto-assign sticker
        while next_sticker in used or next_sticker in skip:
            next_sticker += 1
        sticker_no = next_sticker if next_sticker <= 50 else None
        used.add(next_sticker)
        next_sticker += 1

        item = ThisYearItem(
            year=year,
            sticker_no=sticker_no,
            item_name=item_name,
            category=row.get("category", "").strip(),
            cost=float(row.get("cost", 0) or 0),
            source=row.get("source", "").strip(),
            notes=row.get("notes", "").strip(),
        )
        db.session.add(item)
        count += 1

    db.session.commit()
    return jsonify({"ok": True, "count": count})


@bp.route("/items/export-excel")
@login_required
def pre_items_export_excel():
    """Export this-year items as Excel .xlsx (8 cols)."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    year = int(request.args.get("year", datetime.now().strftime("%Y")))
    items = (
        ThisYearItem.query
        .filter_by(year=year)
        .order_by(ThisYearItem.sticker_no)
        .all()
    )

    wb = Workbook()
    ws = wb.active
    ws.title = f"今年聖物{year}"

    # Title row
    ws.merge_cells("A1:H1")
    ws["A1"] = f"今年聖物 — {year}年 (寶榮堂花炮會)"
    ws["A1"].font = Font(size=14, bold=True, color="8b0000")
    ws["A1"].alignment = Alignment(horizontal="center")

    # Headers
    headers = ["貼紙#", "意頭名", "類別", "成本", "來源", "投得者", "投標金額", "備註"]
    header_fill = PatternFill(start_color="F5ECE6", end_color="F5ECE6", fill_type="solid")
    header_font = Font(bold=True, size=11)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = Alignment(horizontal="center")

    # Data rows
    for i, item in enumerate(items, 4):
        data = [
            item.sticker_no or "",
            item.item_name,
            item.category or "",
            item.cost if item.cost else "",
            item.source or "",
            item.bidder_name or "",
            item.bid_amount if item.bid_amount else "",
            item.notes or "",
        ]
        for col, val in enumerate(data, 1):
            cell = ws.cell(row=i, column=col, value=val)
            cell.border = thin_border
            if col == 1:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal="center")
            if col in (4, 7) and val:
                cell.number_format = "#,##0"
                cell.alignment = Alignment(horizontal="right")

    # Column widths
    widths = [12, 30, 12, 14, 20, 20, 16, 30]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + i)].width = w

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=f"今年聖物_{year}.xlsx",
    )


@bp.route("/items/sticker-pdf")
@login_required
def pre_items_sticker_pdf():
    """Generate sticker labels PDF (A4 3x8 grid, ReportLab, gold accent)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas
    from reportlab.lib.colors import HexColor
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Register CJK font for Chinese character support, fall back to Helvetica if unavailable
    _CJK_PATH = "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"
    try:
        pdfmetrics.registerFont(TTFont("WenQuanYi", _CJK_PATH))
        FONT_NAME = "WenQuanYi"
        FONT_BOLD = "WenQuanYi"
    except Exception:
        FONT_NAME = "Helvetica"
        FONT_BOLD = "Helvetica-Bold"

    year = int(request.args.get("year", datetime.now().strftime("%Y")))
    items = (
        ThisYearItem.query
        .filter_by(year=year)
        .order_by(ThisYearItem.sticker_no)
        .all()
    )

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    # Label dimensions: 3 columns x 8 rows = 24 labels per page
    margin_left = 14 * mm
    margin_top = 20 * mm
    label_w = 63 * mm
    label_h = 35 * mm
    cols, rows = 3, 8
    gold = HexColor("#b8860b")

    if not items:
        # Empty state — show message instead of blank page
        c.setFont(FONT_BOLD, 18)
        c.setFillColor(HexColor("#8c847a"))
        c.drawCentredString(width / 2, height / 2, "暫無聖物資料")
        c.save()
        buf.seek(0)
        return send_file(
            buf,
            mimetype="application/pdf",
            as_attachment=True,
            download_name=f"sticker_labels_{year}.pdf",
        )

    for idx, item in enumerate(items):
        col = idx % cols
        row = (idx // cols) % rows
        page = idx // (cols * rows)

        if page > 0 and col == 0 and row == 0:
            c.showPage()

        x = margin_left + col * label_w
        y = height - margin_top - (row + 1) * label_h

        # Draw label border
        c.setStrokeColor(HexColor("#d1cbc2"))
        c.setLineWidth(0.5)
        c.rect(x, y, label_w - 2 * mm, label_h - 2 * mm)

        # Sticker number (gold accent)
        c.setFillColor(gold)
        c.setFont(FONT_BOLD, 14)
        sticker_no = item.sticker_no if item.sticker_no else "-"
        c.drawString(x + 3 * mm, y + label_h - 10 * mm, f"#{sticker_no}")

        # Item name
        c.setFillColor(HexColor("#2d2822"))
        c.setFont(FONT_NAME, 10)
        name = item.item_name or "-"
        if len(name) > 18:
            name = name[:17] + ".."
        c.drawString(x + 3 * mm, y + label_h - 18 * mm, name)

        # Category
        if item.category:
            c.setFont(FONT_NAME, 8)
            c.setFillColor(HexColor("#8c847a"))
            c.drawString(x + 3 * mm, y + label_h - 25 * mm, item.category)

        # Bid info if won
        if item.bidder_name:
            c.setFont(FONT_BOLD, 8)
            c.setFillColor(HexColor("#dc2626"))
            bid_amt = item.bid_amount or 0
            c.drawString(x + 3 * mm, y + 4 * mm, f"${bid_amt:,.0f}")

    c.save()
    buf.seek(0)

    return send_file(
        buf,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=f"sticker_labels_{year}.pdf",
    )


# ──────────────────────────────────────────────────────────────────────
#  Item Detail API (used by both pre_event and live_event views)
# ──────────────────────────────────────────────────────────────────────


@bp.route("/api/item-detail/<int:item_id>")
@login_required
def item_detail_api(item_id):
    """Return JSON with full item detail including bid info for the detail modal."""
    item = db.session.get(ThisYearItem, item_id)
    if not item:
        return jsonify({"error": "Item not found"}), 404

    is_won = bool(item.bid_amount and item.bid_amount > 0)

    return jsonify({
        "id": item.id,
        "sticker_no": item.sticker_no,
        "item_name": item.item_name or "—",
        "actual_item": item.actual_item or "—",
        "category": item.category or "—",
        "cost": item.cost or 0,
        "source": item.source or "—",
        "notes": item.notes or "",
        "is_won": is_won,
        "bidder_name": item.bidder_name or "",
        "bid_amount": item.bid_amount or 0,
        "paid_amount": item.paid_amount or 0,
        "payment_method": item.payment_method or "",
    })


@bp.route("/api/items/<int:item_id>/detail")
@login_required
def item_detail_from_db(item_id):
    """
    Return JSON with Item detail + bid history for the Item model
    (main items table, used by bid history views).
    """
    item = Item.query.filter_by(item_id=item_id).first()
    if not item:
        return jsonify({"error": "Item not found"}), 404

    item_name = item.name_1_auspicious or ""
    actual_item = item.name_2_description or ""

    # All bids for this item, newest first
    bids = (
        Bid.query
        .filter(Bid.item_id == item_id)
        .order_by(Bid.year.desc(), Bid.bid_no)
        .all()
    )

    bid_records = []
    for b in bids:
        member_name = b.member.name if b.member else ""
        unpaid = (b.bid_amount or 0) - (b.paid_amount or 0)
        bid_records.append({
            "year": b.year,
            "member_name": member_name,
            "bid_amount": b.bid_amount or 0,
            "paid_amount": b.paid_amount or 0,
            "unpaid": max(unpaid, 0),
        })

    is_won = len(bids) > 0
    latest_bid = bids[0] if bids else None

    return jsonify({
        "id": item.item_id,
        "sticker_no": item.item_id,
        "item_name": item_name or "—",
        "actual_item": actual_item or "—",
        "category": "—",
        "cost": 0,
        "source": "—",
        "notes": "",
        "is_won": is_won,
        "bidder_name": latest_bid.member.name if latest_bid and latest_bid.member else "",
        "bid_amount": latest_bid.bid_amount if latest_bid else 0,
        "paid_amount": latest_bid.paid_amount if latest_bid else 0,
        "payment_method": latest_bid.payment_method if latest_bid else "",
        "bids": bid_records,
    })


# ──────────────────────────────────────────────────────────────────────
#  B) 支出管理 — Expenses CRUD
# ──────────────────────────────────────────────────────────────────────

@bp.route("/expenses")
@login_required
def pre_expenses():
    """前期支出 — Pre-event expenses list + total."""
    year = request.args.get("year", datetime.now().strftime("%Y"))
    items = (
        Expense.query
        .filter_by(year=int(year), source="pre")
        .order_by(Expense.id.desc())
        .all()
    )
    total = (
        db.session.query(db.func.coalesce(db.func.sum(Expense.amount), 0))
        .filter(Expense.year == int(year), Expense.source == "pre")
        .scalar()
    )
    cash_sum = (
        db.session.query(db.func.coalesce(db.func.sum(Expense.amount), 0))
        .filter(Expense.year == int(year), Expense.source == "pre", Expense.payment_method == "CASH")
        .scalar()
    )
    chq_sum = (
        db.session.query(db.func.coalesce(db.func.sum(Expense.amount), 0))
        .filter(Expense.year == int(year), Expense.source == "pre", Expense.payment_method == "CHQ")
        .scalar()
    )
    fps_sum = (
        db.session.query(db.func.coalesce(db.func.sum(Expense.amount), 0))
        .filter(Expense.year == int(year), Expense.source == "pre", Expense.payment_method == "轉數快")
        .scalar()
    )
    stats = {"total": total, "cash": cash_sum, "chq": chq_sum, "fps": fps_sum}
    return render_template("pre_event/expenses.html", expenses=items, year=year, stats=stats)


@bp.route("/expenses/add", methods=["POST"])
@login_required
def pre_expenses_add():
    """Add a pre-event expense."""
    year = int(request.form.get("year", datetime.now().strftime("%Y")))
    expense = Expense(
        year=year,
        source="pre",
        date=request.form.get("date", ""),
        subject=request.form.get("subject", ""),
        amount=float(request.form.get("amount", 0) or 0),
        payment_method=request.form.get("payment_method", ""),
        handler=request.form.get("handler", ""),
        details=request.form.get("details", ""),
        category=request.form.get("category", ""),
    )
    db.session.add(expense)
    db.session.commit()
    return jsonify({"ok": True})


@bp.route("/expenses/<int:eid>/edit", methods=["POST"])
@login_required
def pre_expenses_edit(eid):
    """Edit a pre-event expense."""
    expense = db.session.get(Expense, eid)
    if not expense or expense.source != "pre":
        return jsonify({"error": "Expense not found"}), 404
    expense.subject = request.form.get("subject", expense.subject)
    expense.amount = float(request.form.get("amount", 0) or 0)
    expense.payment_method = request.form.get("payment_method", expense.payment_method)
    expense.handler = request.form.get("handler", expense.handler)
    expense.details = request.form.get("details", expense.details)
    expense.date = request.form.get("date", expense.date)
    expense.category = request.form.get("category", expense.category)
    db.session.commit()
    return jsonify({"ok": True})


@bp.route("/expenses/<int:eid>/delete", methods=["POST"])
@login_required
def pre_expenses_delete(eid):
    """Delete a pre-event expense."""
    expense = db.session.get(Expense, eid)
    if not expense or expense.source != "pre":
        return jsonify({"error": "Expense not found"}), 404
    db.session.delete(expense)
    db.session.commit()
    return jsonify({"ok": True})


@bp.route("/expenses/categories")
@login_required
def pre_expenses_categories():
    """Return predefined expense categories as JSON list."""
    return jsonify(PREDEFINED_CATEGORIES)


@bp.route("/expenses/import-categories", methods=["POST"])
@login_required
def pre_expenses_import_categories():
    """Import expense categories from Excel file."""
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file"}), 400
    try:
        from openpyxl import load_workbook
        wb = load_workbook(BytesIO(f.read()))
        ws = wb.active
        cats = []
        for row in ws.iter_rows(min_row=1, max_col=1, values_only=True):
            if row[0] and str(row[0]).strip():
                cats.append(str(row[0]).strip())
        if cats:
            PREDEFINED_CATEGORIES.clear()
            PREDEFINED_CATEGORIES.extend(cats)
        return jsonify({"ok": True, "count": len(cats), "categories": cats})
    except Exception as e:
        return jsonify({"error": str(e)}), 400


@bp.route("/expenses/add-category", methods=["POST"])
@login_required
def pre_expenses_add_category():
    """Add a single expense category."""
    cat = request.form.get("category", "").strip()
    if cat and cat not in PREDEFINED_CATEGORIES:
        PREDEFINED_CATEGORIES.append(cat)
    return jsonify({"ok": True, "categories": PREDEFINED_CATEGORIES})


# ──────────────────────────────────────────────────────────────────────
#  C) 贊助 — Sponsors CRUD
# ──────────────────────────────────────────────────────────────────────

@bp.route("/sponsors")
@login_required
def pre_sponsors():
    """贊助列表 — Sponsors list + total."""
    year = request.args.get("year", datetime.now().strftime("%Y"))
    items = (
        Sponsor.query
        .filter_by(year=int(year))
        .order_by(Sponsor.id.desc())
        .all()
    )
    total = (
        db.session.query(db.func.coalesce(db.func.sum(Sponsor.amount), 0))
        .filter(Sponsor.year == int(year))
        .scalar()
    )
    max_amount = max((item.amount or 0) for item in items) if items else 0
    stats = {"total": total, "count": len(items), "highest": max_amount}
    return render_template("pre_event/sponsors.html", sponsors=items, year=year, stats=stats)


@bp.route("/sponsors/add", methods=["POST"])
@login_required
def pre_sponsors_add():
    """Add a sponsor."""
    year = int(request.form.get("year", datetime.now().strftime("%Y")))
    sponsor = Sponsor(
        year=year,
        sponsor_name=request.form.get("sponsor_name", ""),
        amount=float(request.form.get("amount", 0) or 0),
        item_name=request.form.get("item_name", ""),
        details=request.form.get("details", ""),
    )
    db.session.add(sponsor)
    db.session.commit()
    return jsonify({"ok": True})


@bp.route("/sponsors/<int:sid>/edit", methods=["POST"])
@login_required
def pre_sponsors_edit(sid):
    """Edit a sponsor."""
    sponsor = db.session.get(Sponsor, sid)
    if not sponsor:
        return jsonify({"error": "Sponsor not found"}), 404
    sponsor.sponsor_name = request.form.get("sponsor_name", sponsor.sponsor_name)
    sponsor.amount = float(request.form.get("amount", 0) or 0)
    sponsor.item_name = request.form.get("item_name", sponsor.item_name)
    sponsor.details = request.form.get("details", sponsor.details)
    db.session.commit()
    return jsonify({"ok": True})


@bp.route("/sponsors/<int:sid>/delete", methods=["POST"])
@login_required
def pre_sponsors_delete(sid):
    """Delete a sponsor."""
    sponsor = db.session.get(Sponsor, sid)
    if not sponsor:
        return jsonify({"error": "Sponsor not found"}), 404
    db.session.delete(sponsor)
    db.session.commit()
    return jsonify({"ok": True})


# ──────────────────────────────────────────────────────────────────────
#  D) 上年欠款 / 粉紅紙 / 收據
# ──────────────────────────────────────────────────────────────────────

@bp.route("/previous")
@login_required
def pre_previous():
    """上年欠款 — List members with unpaid from previous years."""
    year = request.args.get("year", "").strip()
    search = request.args.get("q", "").strip()

    # Build base query — members who have bids with bid_amount > 0
    subq = (
        db.session.query(Bid.member_id)
        .filter(Bid.bid_amount > 0)
        .distinct()
        .subquery()
    )

    query = db.session.query(
        Member.member_id,
        Member.name,
        Member.phone,
        db.func.coalesce(
            db.session.query(db.func.sum(Bid.bid_amount))
            .filter(Bid.member_id == Member.member_id, Bid.bid_amount > 0)
            .scalar_subquery(), 0
        ).label("total_due"),
        db.func.coalesce(
            db.session.query(db.func.sum(Bid.paid_amount))
            .filter(Bid.member_id == Member.member_id)
            .scalar_subquery(), 0
        ).label("total_paid"),
        db.func.count(
            db.session.query(Bid.id)
            .filter(Bid.member_id == Member.member_id, Bid.bid_amount > 0)
            .scalar_subquery()
        ).label("item_count"),
    ).filter(Member.member_id.in_(subq))

    # If year is specified, scope to that year
    if year:
        total_due_sub = (
            db.session.query(db.func.coalesce(db.func.sum(Bid.bid_amount), 0))
            .filter(Bid.member_id == Member.member_id, Bid.year == int(year), Bid.bid_amount > 0)
            .scalar_subquery()
        )
        total_paid_sub = (
            db.session.query(db.func.coalesce(db.func.sum(Bid.paid_amount), 0))
            .filter(Bid.member_id == Member.member_id, Bid.year == int(year))
            .scalar_subquery()
        )
        item_count_sub = (
            db.session.query(db.func.count(Bid.id))
            .filter(Bid.member_id == Member.member_id, Bid.year == int(year), Bid.bid_amount > 0)
            .scalar_subquery()
        )
        query = (
            db.session.query(
                Member.member_id,
                Member.name,
                Member.phone,
                total_due_sub.label("total_due"),
                total_paid_sub.label("total_paid"),
                item_count_sub.label("item_count"),
            )
            .filter(Member.member_id.in_(
                db.session.query(Bid.member_id)
                .filter(Bid.year == int(year), Bid.bid_amount > 0)
                .distinct()
            ))
        )

    rows = query.all()
    result = []
    for r in rows:
        d = r._asdict()
        d["unpaid"] = (d["total_due"] or 0) - (d["total_paid"] or 0)
        d["member_name"] = d.get("name", "")
        result.append(d)

    # Search filter
    if search:
        name_match_ids = {
            r["member_id"] for r in result
            if search.lower() in (r["name"] or "").lower()
               or search.lower() in (r.get("phone") or "").lower()
        }
        # Also search by item name in bids
        if year:
            item_matches = (
                db.session.query(Bid.member_id)
                .join(Bid.item, isouter=True)
                .filter(
                    Bid.year == int(year),
                    db.or_(
                        Bid.item.has(db.or_(
                            Bid.item.name_1_auspicious.ilike(f"%{search}%"),
                            Bid.item.name_2_description.ilike(f"%{search}%"),
                        ))
                    )
                )
                .distinct()
                .all()
            )
        else:
            item_matches = (
                db.session.query(Bid.member_id)
                .join(Bid.item, isouter=True)
                .filter(
                    db.or_(
                        Bid.item.has(db.or_(
                            Bid.item.name_1_auspicious.ilike(f"%{search}%"),
                            Bid.item.name_2_description.ilike(f"%{search}%"),
                        ))
                    )
                )
                .distinct()
                .all()
            )
        item_match_ids = {r[0] for r in item_matches}
        result = [
            r for r in result
            if r["member_id"] in name_match_ids or r["member_id"] in item_match_ids
        ]

    years = [r[0] for r in db.session.query(Bid.year).distinct().order_by(Bid.year.desc()).all()]

    # Compute stats for template
    total_due = sum((r.get("total_due") or 0) for r in result)
    total_paid = sum((r.get("total_paid") or 0) for r in result)
    total_unpaid = sum((r.get("unpaid") or 0) for r in result)
    unpaid_count = sum(1 for r in result if (r.get("unpaid") or 0) > 0)

    return render_template(
        "pre_event/previous.html",
        debts=result,
        years=years, year=year, search=search,
        stats={
            "total_due": total_due,
            "total_paid": total_paid,
            "unpaid": total_unpaid,
            "unpaid_count": unpaid_count,
        },
    )


@bp.route("/previous/<int:member_id>")
@login_required
def pre_previous_member(member_id):
    """上年會員欠款詳情 — Member detail with unpaid items."""
    year = request.args.get("year", datetime.now().strftime("%Y"))
    member = Member.query.filter_by(member_id=member_id).first()
    if not member:
        return "Member not found", 404

    bids = (
        db.session.query(Bid)
        .filter(Bid.member_id == member_id, Bid.year == int(year))
        .order_by(Bid.bid_no)
        .all()
    )

    total = sum(b.bid_amount or 0 for b in bids)
    paid = sum(b.paid_amount or 0 for b in bids)

    # Add computed fields to member for template
    member.total_due = total
    member.total_paid = paid
    member.unpaid_total = total - paid
    member.item_count = len(bids)

    return render_template(
        "pre_event/previous_member.html",
        member=member, items=bids,
        year=year, total=total, paid=paid,
        unpaid=total - paid,
    )


@bp.route("/invoice/<int:member_id>")
@login_required
def generate_invoice(member_id):
    """Generate 粉紅紙 (Pink Slip) DOCX — with template fallback."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from pathlib import Path

    year = int(request.args.get("year", datetime.now().strftime("%Y")))
    member = Member.query.filter_by(member_id=member_id).first()
    if not member:
        return "Member not found", 404

    bids = (
        db.session.query(Bid)
        .filter(Bid.member_id == member_id, Bid.year == year)
        .order_by(Bid.bid_no)
        .all()
    )

    total = sum(b.bid_amount or 0 for b in bids)
    fee = sum(b.membership_fee or 0 for b in bids)

    from flask import current_app
    uploads = Path(current_app.config.get("UPLOADS_DIR", "uploads"))
    template_path = uploads / f"粉紅紙{year}-template.docx"

    if template_path.exists():
        doc = Document(str(template_path))
        if doc.tables:
            t = doc.tables[0]
            # Add item rows dynamically
            for _ in range(len(bids)):
                t.add_row()
            for idx, bid in enumerate(bids):
                row_idx = 1 + idx
                if row_idx < len(t.rows):
                    cells = t.rows[row_idx].cells
                else:
                    cells = t.rows[-1].cells
                if len(cells) >= 3:
                    cells[0].text = str(idx + 1)
                    desc = bid.item.name_1_auspicious if bid.item and bid.item.name_1_auspicious else \
                           bid.item.name_2_description if bid.item else "-"
                    cells[1].text = desc if desc else "-"
                    cells[min(2, len(cells) - 1)].text = f"${bid.bid_amount or 0:,.0f}"
    else:
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("🏮 寶榮堂花炮會 — 粉紅紙")
        run.bold = True
        run.font.size = Pt(18)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Pink Slip — {year}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(f"會員: {member.name} (編號: {member.member_id})")
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph()
        table = doc.add_table(rows=1 + len(bids) + 1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, txt in enumerate(["#", "聖物名稱", "金額"]):
            hdr[i].text = txt
        for idx, bid in enumerate(bids):
            row = table.rows[1 + idx].cells
            row[0].text = str(idx + 1)
            desc = bid.item.name_1_auspicious if bid.item and bid.item.name_1_auspicious else \
                   bid.item.name_2_description if bid.item else "-"
            row[1].text = desc if desc else "-"
            row[2].text = f"${bid.bid_amount or 0:,.0f}"
        total_row = table.rows[-1].cells
        total_row[0].text = ""
        total_row[1].text = "總計:"
        total_row[2].text = f"${total:,.0f}"

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"粉紅紙_{member.name}_{year}.docx",
    )


@bp.route("/receipt/<int:member_id>")
@login_required
def generate_receipt(member_id):
    """Generate PWT RECEIPT XLSX — 3 receipts per A4 (正本/副本/存根)."""
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Font, Alignment, Border, Side
    from pathlib import Path

    year = int(request.args.get("year", datetime.now().strftime("%Y")))
    member = Member.query.filter_by(member_id=member_id).first()
    if not member:
        return "Member not found", 404

    bids = (
        db.session.query(Bid)
        .filter(Bid.member_id == member_id, Bid.year == year)
        .order_by(Bid.bid_no)
        .all()
    )

    total = sum(b.bid_amount or 0 for b in bids)

    from flask import current_app
    uploads = Path(current_app.config.get("UPLOADS_DIR", "uploads"))
    template_path = uploads / f"PWT RECEIPT {year}-template.xlsx"

    if template_path.exists():
        wb = load_workbook(str(template_path))
        ws = wb.active
    else:
        wb = Workbook()
        ws = wb.active

    ws.title = "Receipt"
    ws.page_setup.orientation = "portrait"
    ws.page_setup.paperSize = ws.PAPERSIZE_A4

    thin = Side(style="thin")
    bold_font = Font(bold=True, size=11)
    normal_font = Font(size=10)
    title_font = Font(bold=True, size=14)

    def draw_receipt(start_row, copy_label):
        """Draw one receipt block starting at start_row."""
        r = start_row
        ws.cell(row=r, column=1, value="寶榮堂花炮會").font = title_font
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
        r += 1
        ws.cell(row=r, column=1, value=f"PWT RECEIPT — {copy_label}").font = Font(size=10, color="666666")
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
        r += 1
        ws.cell(row=r, column=1, value="會員編號:").font = bold_font
        ws.cell(row=r, column=2, value=member.member_id).font = normal_font
        ws.cell(row=r, column=4, value="日期:").font = bold_font
        ws.cell(row=r, column=5, value=f"{year}年").font = normal_font
        r += 1
        ws.cell(row=r, column=1, value="會員姓名:").font = bold_font
        ws.cell(row=r, column=2, value=member.name).font = Font(bold=True, size=12)
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=4)
        r += 1
        for ci, h in enumerate(["#", "聖物名稱", "金額"], 1):
            c = ws.cell(row=r, column=ci, value=h)
            c.font = bold_font
            c.border = Border(top=thin, bottom=thin)
        r += 1
        for idx, bid in enumerate(bids):
            ws.cell(row=r, column=1, value=idx + 1).font = normal_font
            desc = bid.item.name_1_auspicious if bid.item and bid.item.name_1_auspicious else \
                   bid.item.name_2_description if bid.item else "-"
            ws.cell(row=r, column=2, value=desc if desc else "-").font = normal_font
            ws.cell(row=r, column=3, value=bid.bid_amount or 0).font = normal_font
            ws.cell(row=r, column=3).number_format = "$#,##0"
            r += 1
        ws.cell(row=r, column=1, value="").font = normal_font
        ws.cell(row=r, column=2, value="總計:").font = bold_font
        ws.cell(row=r, column=3, value=total).font = Font(bold=True, size=12)
        ws.cell(row=r, column=3).number_format = "$#,##0"
        r += 2
        return r

    next_row = draw_receipt(1, "正本")
    next_row = draw_receipt(next_row + 1, "副本")
    next_row = draw_receipt(next_row + 1, "存根")

    ws.column_dimensions["A"].width = 14
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 16
    ws.column_dimensions["D"].width = 12
    ws.column_dimensions["E"].width = 16
    ws.column_dimensions["F"].width = 8

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=f"RECEIPT_{member.name}_{year}.xlsx",
    )


@bp.route("/previous/batch-pdf", methods=["POST"])
@login_required
def pre_previous_batch_pdf():
    """Generate a combined PDF for selected members (batch invoice)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas
    from reportlab.lib.colors import HexColor

    year = int(request.form.get("year", datetime.now().strftime("%Y")))
    member_ids = request.form.getlist("member_ids")
    if not member_ids:
        # Support JSON body
        data = request.get_json(silent=True)
        if data and "member_ids" in data:
            member_ids = data["member_ids"]

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    for mid in member_ids:
        try:
            mid = int(mid)
        except (ValueError, TypeError):
            continue

        member = Member.query.filter_by(member_id=mid).first()
        if not member:
            continue

        bids = (
            db.session.query(Bid)
            .filter(Bid.member_id == mid, Bid.year == year)
            .order_by(Bid.bid_no)
            .all()
        )
        if not bids:
            continue

        total = sum(b.bid_amount or 0 for b in bids)

        # Title
        c.setFont("Helvetica-Bold", 16)
        c.setFillColor(HexColor("#8b0000"))
        c.drawCentredString(width / 2, height - 30 * mm, "寶榮堂花炮會 — 粉紅紙")
        c.setFont("Helvetica", 10)
        c.setFillColor(HexColor("#666666"))
        c.drawCentredString(width / 2, height - 36 * mm, f"Pink Slip — {year}")

        # Member info
        c.setFillColor(HexColor("#000000"))
        c.setFont("Helvetica-Bold", 12)
        c.drawString(20 * mm, height - 48 * mm, f"會員: {member.name} (編號: {member.member_id})")

        # Table header
        y_start = height - 58 * mm
        row_h = 7 * mm
        col_x = [20 * mm, 40 * mm, 140 * mm]
        c.setFont("Helvetica-Bold", 10)
        c.drawString(col_x[0], y_start, "#")
        c.drawString(col_x[1], y_start, "聖物名稱")
        c.drawString(col_x[2], y_start, "金額")

        y = y_start - row_h
        for idx, bid in enumerate(bids):
            c.setFont("Helvetica", 10)
            c.drawString(col_x[0], y, str(idx + 1))
            desc = bid.item.name_1_auspicious if bid.item and bid.item.name_1_auspicious else \
                   bid.item.name_2_description if bid.item else "-"
            c.drawString(col_x[1], y, (desc or "-")[:30])
            c.drawString(col_x[2], y, f"${bid.bid_amount or 0:,.0f}")
            y -= row_h

        # Total
        c.setFont("Helvetica-Bold", 11)
        c.drawString(col_x[1], y, "總計:")
        c.drawString(col_x[2], y, f"${total:,.0f}")

        c.showPage()

    c.save()
    buf.seek(0)

    return send_file(
        buf,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=f"batch_invoice_{year}.pdf",
    )


@bp.route("/previous/batch-export", methods=["POST"])
@login_required
def pre_previous_batch_export():
    """Batch export DOCX invoices as ZIP."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from pathlib import Path

    year = int(request.form.get("year", datetime.now().strftime("%Y")))
    member_ids = request.form.getlist("member_ids")
    if not member_ids:
        data = request.get_json(silent=True)
        if data and "member_ids" in data:
            member_ids = data["member_ids"]

    zip_buf = BytesIO()

    with ZipFile(zip_buf, "w") as zf:
        for mid in member_ids:
            try:
                mid = int(mid)
            except (ValueError, TypeError):
                continue

            member = Member.query.filter_by(member_id=mid).first()
            if not member:
                continue

            bids = (
                db.session.query(Bid)
                .filter(Bid.member_id == mid, Bid.year == year)
                .order_by(Bid.bid_no)
                .all()
            )
            if not bids:
                continue

            total = sum(b.bid_amount or 0 for b in bids)

            doc = Document()
            for section in doc.sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(1.5)
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("🏮 寶榮堂花炮會 — 粉紅紙")
            run.bold = True
            run.font.size = Pt(18)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Pink Slip — {year}")
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(100, 100, 100)

            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(f"會員: {member.name} (編號: {member.member_id})")
            run.bold = True
            run.font.size = Pt(12)

            doc.add_paragraph()
            table = doc.add_table(rows=1 + len(bids) + 1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for i, txt in enumerate(["#", "聖物名稱", "金額"]):
                hdr[i].text = txt
            for idx, bid in enumerate(bids):
                row = table.rows[1 + idx].cells
                row[0].text = str(idx + 1)
                desc = bid.item.name_1_auspicious if bid.item and bid.item.name_1_auspicious else \
                       bid.item.name_2_description if bid.item else "-"
                row[1].text = desc if desc else "-"
                row[2].text = f"${bid.bid_amount or 0:,.0f}"
            total_row = table.rows[-1].cells
            total_row[0].text = ""
            total_row[1].text = "總計:"
            total_row[2].text = f"${total:,.0f}"

            docx_buf = BytesIO()
            doc.save(docx_buf)
            docx_buf.seek(0)
            safe_name = member.name.replace("/", "_").replace("\\", "_")
            zf.writestr(f"粉紅紙_{safe_name}_{year}.docx", docx_buf.getvalue())

    zip_buf.seek(0)
    return send_file(
        zip_buf,
        mimetype="application/zip",
        as_attachment=True,
        download_name=f"batch_invoice_{year}.zip",
    )
